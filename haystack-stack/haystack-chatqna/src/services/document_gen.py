"""
Document Generation Service — generates PDF + DOCX from conversation history.

Pipeline:
  1. LLM reads conversation → generates structured document content (JSON)
  2. Renderer builds PDF or DOCX from the structured content
  3. Returns bytes + metadata for frontend preview/download

Document types:
  - consultation_summary: what was discussed, recommendations, follow-up
  - care_plan: personalised care plan document
  - referral_letter: for facility referral
  - health_report: general health status report
"""

from typing import Dict, Any, Optional, List
from datetime import datetime
import json
import io
import logging
import uuid

from openai import AsyncOpenAI
from src.config import settings

logger = logging.getLogger(__name__)


# ── i18n: hardcoded labels used by the PDF/DOCX renderers.
# The LLM-generated document content (title, sections, recommendations,
# etc.) is translated per-field in translate_doc_for_render().
# The static chrome strings below — section headers, disclaimers,
# cover-block labels — are NOT LLM output, so they're kept as a small
# local catalogue. When language != "en", the route handler translates
# them via the existing Translator service before passing to render.
DEFAULT_LABELS = {
    # structured-document renderer (render_pdf / render_docx)
    "recommendations":  "Recommendations",
    "follow_up":        "Follow-up Plan",
    "generated":        "Generated",

    # conversation-transcript renderer (render_conversation_pdf)
    "transcript_header":   "AMINA CARE  |  Conversation Transcript",
    "transcript_title":    "Conversation Transcript",
    "prepared_for":        "Prepared for",
    "session_date":        "Session date",
    "turns_recorded":      "Turns recorded",
    "no_conversation":     "No conversation to export yet. Chat with AMINA first.",
    "care_agent_label":    "AMINA Care Agent",
    "patient_label":       "Patient",
    "doc_id":              "Doc ID",
    "page_of":             "Page {current} of {total}",
    "confidential_footer": "AMINA Care  -  Confidential patient document",
    "disclaimer_heading":  "Disclaimer",
    "source_prefix":       "Source",
    "transcript_disclaimer": (
        "This transcript is a verbatim record of the conversation between the "
        "patient and AMINA, an AI-assisted community health agent. It is intended "
        "as an informational record only and does not constitute a medical "
        "diagnosis, treatment plan, or substitute for advice from a qualified "
        "healthcare professional. Any clinical recommendations should be reviewed "
        "with a licensed provider. AMINA Care holds this document as confidential "
        "patient data; unauthorised distribution is prohibited."
    ),
}


async def _translate_labels(language: str) -> Dict[str, str]:
    """Translate the DEFAULT_LABELS catalogue for a target language.

    Returns English labels unchanged when language == "en" or when the
    translator is unavailable (safe fallback). Uses the existing batch
    translator so cache hits are shared across requests.
    """
    if not language or language == "en":
        return dict(DEFAULT_LABELS)
    try:
        from src.services.translator import get_translator
        tr = get_translator()
        translated = await tr.translate_batch(
            DEFAULT_LABELS,
            source="en",
            target=language,
        )
        # translate_batch returns {key: translated_text}; merge over defaults
        # so any missing key falls back to English rather than empty string.
        merged = dict(DEFAULT_LABELS)
        for k, v in (translated or {}).items():
            if v:
                merged[k] = v
        return merged
    except Exception as e:
        logger.warning(f"Label translation failed, using English: {e}")
        return dict(DEFAULT_LABELS)


async def translate_doc_for_render(
    doc: Dict[str, Any],
    language: str,
) -> Dict[str, Any]:
    """Translate a structured-document dict's user-visible fields.

    Fields translated: title, subtitle, sections[*].heading/content,
    recommendations[*], follow_up, disclaimer.
    Metadata fields (doc_id, doc_type, generated_at, session_id) are
    preserved verbatim.

    Returns the original dict unchanged if language == "en" or on failure.
    """
    if not doc or not language or language == "en":
        return doc
    try:
        from src.services.translator import get_translator
        tr = get_translator()
    except Exception as e:
        logger.warning(f"Translator unavailable, returning English doc: {e}")
        return doc

    try:
        flat = {}
        if doc.get("title"):      flat["title"]      = doc["title"]
        if doc.get("subtitle"):   flat["subtitle"]   = doc["subtitle"]
        if doc.get("follow_up"):  flat["follow_up"]  = doc["follow_up"]
        if doc.get("disclaimer"): flat["disclaimer"] = doc["disclaimer"]
        for i, s in enumerate(doc.get("sections", []) or []):
            if s.get("heading"):  flat[f"sec{i}_heading"] = s["heading"]
            if s.get("content"):  flat[f"sec{i}_content"] = s["content"]
        for i, r in enumerate(doc.get("recommendations", []) or []):
            if r: flat[f"rec{i}"] = r

        translated = await tr.translate_batch(flat, source="en", target=language)

        out = dict(doc)
        if "title" in translated:      out["title"]      = translated["title"]
        if "subtitle" in translated:   out["subtitle"]   = translated["subtitle"]
        if "follow_up" in translated:  out["follow_up"]  = translated["follow_up"]
        if "disclaimer" in translated: out["disclaimer"] = translated["disclaimer"]

        new_sections = []
        for i, s in enumerate(doc.get("sections", []) or []):
            ns = dict(s)
            if f"sec{i}_heading" in translated: ns["heading"] = translated[f"sec{i}_heading"]
            if f"sec{i}_content" in translated: ns["content"] = translated[f"sec{i}_content"]
            new_sections.append(ns)
        if new_sections:
            out["sections"] = new_sections

        new_recs = []
        for i, r in enumerate(doc.get("recommendations", []) or []):
            new_recs.append(translated.get(f"rec{i}", r))
        if new_recs:
            out["recommendations"] = new_recs

        return out
    except Exception as e:
        logger.warning(f"Doc translation failed, returning English: {e}")
        return doc


async def translate_messages_for_render(
    messages: List[Dict[str, str]],
    language: str,
) -> List[Dict[str, str]]:
    """Translate the `content` of every conversation message into the
    target language. Roles and sources are preserved verbatim.

    Returns the original list unchanged if language == "en" or on failure.
    """
    if not messages or not language or language == "en":
        return messages
    try:
        from src.services.translator import get_translator
        tr = get_translator()
    except Exception as e:
        logger.warning(f"Translator unavailable, messages stay English: {e}")
        return messages

    try:
        flat = {}
        for i, m in enumerate(messages):
            c = (m or {}).get("content") or ""
            if c.strip():
                flat[f"m{i}"] = c
        translated = await tr.translate_batch(flat, source="en", target=language)

        out = []
        for i, m in enumerate(messages):
            nm = dict(m or {})
            if f"m{i}" in translated and translated[f"m{i}"]:
                nm["content"] = translated[f"m{i}"]
            out.append(nm)
        return out
    except Exception as e:
        logger.warning(f"Message translation failed, kept English: {e}")
        return messages


DOCUMENT_GENERATION_PROMPT = """You are creating a formal health document for a Gambian patient based on their
conversation with AMINA (a community health worker AI).

Read the conversation below and generate a structured document.

Document type: {doc_type}

CONVERSATION:
{transcript}

Patient context:
- Name: {patient_name}
- Conditions: {conditions}
- Region: {region}

Generate the document as JSON with this EXACT structure:
{{
  "title": "Document title",
  "subtitle": "Patient name + date",
  "sections": [
    {{
      "heading": "Section heading",
      "content": "Paragraph text for this section. Use clear, professional language."
    }}
  ],
  "recommendations": ["recommendation 1", "recommendation 2", ...],
  "follow_up": "Next steps and follow-up plan",
  "disclaimer": "This document is generated by AMINA AI health assistant and should be reviewed by a qualified healthcare provider."
}}

For a {doc_type}, include these sections:
- Patient Information (name, age, conditions, region)
- Consultation Summary (what was discussed)
- Clinical Findings (vitals, symptoms, assessments)
- Recommendations (medications, diet, lifestyle)
- Follow-up Plan (next check-in, warning signs)
- References (WHO PEN guidelines used)

Use professional medical document language but keep it simple enough for the patient.
Return ONLY valid JSON, no markdown fences.
"""


async def generate_document_content(
    messages: List[Dict[str, str]],
    patient_name: str = "Patient",
    conditions: List[str] = None,
    region: str = "Kanifing",
    doc_type: str = "consultation_summary",
) -> Dict[str, Any]:
    """Use LLM to generate structured document content from conversation."""
    client = AsyncOpenAI(
        api_key=settings.OPENAI_API_KEY,
        base_url=settings.OPENAI_BASE_URL,
    )

    transcript = "\n".join(
        f"{m.get('role', 'user').upper()}: {m.get('content', '')}"
        for m in messages[-20:]
    )

    prompt = DOCUMENT_GENERATION_PROMPT.format(
        doc_type=doc_type,
        transcript=transcript,
        patient_name=patient_name,
        conditions=", ".join(conditions) if conditions else "not specified",
        region=region,
    )

    try:
        resp = await client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1200,
        )
        text = (resp.choices[0].message.content or "").strip()
        if text.startswith("```"):
            import re
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"```\s*$", "", text)
        return json.loads(text)
    except Exception as e:
        logger.error(f"Document generation failed: {e}")
        return {
            "title": f"Health {doc_type.replace('_', ' ').title()}",
            "subtitle": f"{patient_name} — {datetime.now().strftime('%d %B %Y')}",
            "sections": [
                {"heading": "Summary", "content": "Document generation encountered an error. Please try again."}
            ],
            "recommendations": [],
            "follow_up": "",
            "disclaimer": "Generated by AMINA AI.",
        }


def render_pdf(doc: Dict[str, Any], labels: Optional[Dict[str, str]] = None) -> bytes:
    """Render a structured document to PDF bytes.

    `labels` overrides the hardcoded chrome strings (section headers,
    "Generated:" etc.) so the document renders in a non-English language
    when the route handler pre-translates via _translate_labels().
    When None, English defaults are used.
    """
    from fpdf import FPDF

    L = {**DEFAULT_LABELS, **(labels or {})}

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, doc.get("title", "Health Document"), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, doc.get("subtitle", ""), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(5)

    # Date
    pdf.set_font("Helvetica", "I", 9)
    pdf.cell(0, 6, f"{L['generated']}: {datetime.now().strftime('%d %B %Y, %H:%M')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    # Sections
    for section in doc.get("sections", []):
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, section.get("heading", ""), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        content = section.get("content", "")
        pdf.multi_cell(0, 6, content)
        pdf.ln(3)

    # Recommendations
    recs = doc.get("recommendations", [])
    if recs:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, L["recommendations"], new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for i, rec in enumerate(recs, 1):
            try:
                pdf.multi_cell(0, 6, "{}. {}".format(i, rec))
            except Exception:
                pdf.cell(0, 6, "{}. {}".format(i, rec[:80]), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    # Follow-up
    fu = doc.get("follow_up", "")
    if fu:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, L["follow_up"], new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, fu)
        pdf.ln(3)

    # Disclaimer
    pdf.ln(8)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(3)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(0, 5, doc.get("disclaimer", "Generated by AMINA AI."))

    return bytes(pdf.output())


def _conv_sanitize(text: str) -> str:
    """Clean text for fpdf2 rendering. Preserves UTF-8 for Unicode fonts."""
    if text is None:
        return ""
    return (
        str(text)
        .replace("\u2019", "'").replace("\u2018", "'")
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2013", "-").replace("\u2014", "-")
        .replace("\u2026", "...")
        .replace("\u00a0", " ")
    )


def render_conversation_pdf(
    messages: List[Dict[str, str]],
    patient_name: str = "Patient",
    language: str = "en",
    labels: Optional[Dict[str, str]] = None,
) -> bytes:
    """Render a raw chat conversation transcript to PDF bytes.

    Business-style layout with branded header, per-page footer, page
    numbering, and a structured message block per turn. Unlike
    render_pdf() (which uses the LLM to produce a structured document),
    this emits a verbatim transcript of the user ↔ AMINA session.

    `messages` should already be pre-translated into `language` by the
    caller (via translate_messages_for_render) — this function only
    renders. `labels` overrides chrome strings to match `language`.
    """
    from fpdf import FPDF

    L = {**DEFAULT_LABELS, **(labels or {})}

    doc_id      = "CHAT-{}".format(uuid.uuid4().hex[:8].upper())
    generated   = datetime.now().strftime("%d %B %Y, %H:%M")
    safe_name   = _conv_sanitize(patient_name or L["patient_label"])
    header_text = _conv_sanitize(L["transcript_header"])
    page_w      = 210  # A4 portrait mm
    margin      = 15
    doc_id_label = _conv_sanitize(L["doc_id"])
    generated_label = _conv_sanitize(L["generated"])
    confidential = _conv_sanitize(L["confidential_footer"])
    page_of_tmpl = L["page_of"]  # "Page {current} of {total}"

    use_unicode = language and language != "en"
    _font = "dejavu" if use_unicode else "Helvetica"

    class AminaTranscriptPDF(FPDF):
        def _setup_fonts(self):
            if use_unicode:
                import os, glob
                search = [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                    "/usr/share/fonts/TTF/DejaVuSans.ttf",
                    os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans.ttf"),
                ]
                search += glob.glob("/usr/share/fonts/**/DejaVuSans.ttf", recursive=True)
                ttf = None
                for p in search:
                    if os.path.isfile(p):
                        ttf = p
                        break
                if ttf:
                    self.add_font("dejavu", "", ttf, uni=True)
                    self.add_font("dejavu", "B", ttf.replace("Sans.", "Sans-Bold.") if os.path.isfile(ttf.replace("Sans.", "Sans-Bold.")) else ttf, uni=True)
                    self.add_font("dejavu", "I", ttf.replace("Sans.", "Sans-Oblique.") if os.path.isfile(ttf.replace("Sans.", "Sans-Oblique.")) else ttf, uni=True)
                else:
                    logger.warning("DejaVuSans.ttf not found, falling back to Helvetica")
                    nonlocal _font
                    _font = "Helvetica"

        def header(self):  # noqa: D401
            self.set_fill_color(12, 17, 40)         # surface
            self.rect(0, 0, page_w, 18, style="F")
            self.set_text_color(226, 232, 240)      # text
            self.set_font(_font, "B", 12)
            self.set_xy(margin, 5)
            self.cell(page_w - 2 * margin, 8, header_text, align="L")
            self.set_font(_font, "", 9)
            self.set_xy(margin, 11)
            self.set_text_color(148, 163, 184)      # muted
            self.cell(
                page_w - 2 * margin, 5,
                "{}: {}    {}: {}".format(doc_id_label, doc_id, generated_label, generated),
                align="L",
            )
            # Accent bar
            self.set_fill_color(129, 140, 248)      # accent indigo
            self.rect(0, 18, page_w, 1.2, style="F")
            self.set_text_color(0, 0, 0)
            self.set_y(24)

        def footer(self):
            self.set_y(-14)
            self.set_draw_color(200, 200, 210)
            self.line(margin, self.get_y(), page_w - margin, self.get_y())
            self.set_y(-11)
            self.set_font(_font, "I", 8)
            self.set_text_color(120, 120, 130)
            left = confidential
            # fpdf renders {nb} lazily; use placeholder for total, current is
            # known at call time. The template may contain {current} / {total}
            # in the translated string — keep both so translators don't have
            # to preserve exact tokens. Fall back to English layout if the
            # translation dropped either placeholder.
            try:
                right = _conv_sanitize(
                    page_of_tmpl.format(current=self.page_no(), total="{nb}")
                )
            except Exception:
                right = "Page {} of {{nb}}".format(self.page_no())
            self.cell(0, 5, left, align="L")
            self.set_y(-11)
            self.cell(0, 5, right, align="R")
            self.set_text_color(0, 0, 0)

    pdf = AminaTranscriptPDF(orientation="P", unit="mm", format="A4")
    if use_unicode:
        pdf._setup_fonts()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_left_margin(margin)
    pdf.set_right_margin(margin)
    pdf.add_page()

    # ── Cover block ─────────────────────────────────────────────────────────
    pdf.ln(4)
    pdf.set_font(_font, "B", 17)
    pdf.set_text_color(15, 23, 42)
    pdf.cell(0, 9, _conv_sanitize(L["transcript_title"]), new_x="LMARGIN", new_y="NEXT")

    pdf.set_font(_font, "", 10)
    pdf.set_text_color(71, 85, 105)
    pdf.cell(0, 6, "{}: {}".format(_conv_sanitize(L["prepared_for"]), safe_name), new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, "{}: {}".format(_conv_sanitize(L["session_date"]), generated), new_x="LMARGIN", new_y="NEXT")
    pdf.cell(
        0, 6,
        "{}: {}".format(
            _conv_sanitize(L["turns_recorded"]),
            len([m for m in (messages or []) if m.get("content")]),
        ),
        new_x="LMARGIN", new_y="NEXT",
    )
    pdf.ln(3)
    pdf.set_draw_color(226, 232, 240)
    pdf.line(margin, pdf.get_y(), page_w - margin, pdf.get_y())
    pdf.ln(5)

    if not messages:
        pdf.set_font(_font, "I", 11)
        pdf.set_text_color(120, 120, 130)
        pdf.multi_cell(0, 6, _conv_sanitize(L["no_conversation"]))
        return bytes(pdf.output())

    # ── Message blocks ──────────────────────────────────────────────────────
    for idx, m in enumerate(messages, 1):
        role = (m.get("role") or "user").lower()
        if role == "assistant":
            label       = _conv_sanitize(L["care_agent_label"])
            label_rgb   = (79, 70, 229)   # indigo-600
            band_rgb    = (238, 242, 255) # indigo-50
        elif role == "user":
            label       = "{} ({})".format(_conv_sanitize(L["patient_label"]), safe_name)
            label_rgb   = (15, 23, 42)    # slate-900
            band_rgb    = (248, 250, 252) # slate-50
        else:
            label       = role.upper()
            label_rgb   = (100, 116, 139)
            band_rgb    = (248, 250, 252)

        # Role band
        pdf.set_fill_color(*band_rgb)
        pdf.set_draw_color(226, 232, 240)
        start_y = pdf.get_y()
        pdf.rect(margin, start_y, page_w - 2 * margin, 7, style="FD")

        pdf.set_xy(margin + 2, start_y + 0.8)
        pdf.set_font(_font, "B", 9)
        pdf.set_text_color(*label_rgb)
        pdf.cell((page_w - 2 * margin) / 2, 5.5, "#{:02d}  {}".format(idx, label), align="L")

        pdf.set_xy(margin, start_y + 0.8)
        pdf.set_font(_font, "I", 8)
        pdf.set_text_color(120, 120, 130)
        pdf.cell(page_w - 2 * margin - 2, 5.5, role.capitalize(), align="R")

        pdf.set_y(start_y + 7)
        pdf.ln(2)

        # Body
        pdf.set_font(_font, "", 10)
        pdf.set_text_color(30, 41, 59)
        content = _conv_sanitize(m.get("content") or "")
        try:
            pdf.multi_cell(0, 5.6, content)
        except Exception:
            pdf.multi_cell(0, 5.6, content[:1500])

        # Sources (assistant only)
        sources = m.get("sources") or []
        if sources and isinstance(sources, list):
            pdf.ln(1)
            pdf.set_font(_font, "I", 8)
            pdf.set_text_color(100, 116, 139)
            for s in sources[:3]:
                if isinstance(s, dict):
                    src_text = _conv_sanitize(s.get("title") or s.get("url") or "")
                else:
                    src_text = _conv_sanitize(s)
                if src_text:
                    pdf.multi_cell(0, 4.6, "     {}: {}".format(_conv_sanitize(L["source_prefix"]), src_text[:140]))
            pdf.set_text_color(30, 41, 59)

        pdf.ln(4)

    # ── Footer disclaimer ───────────────────────────────────────────────────
    pdf.ln(4)
    pdf.set_draw_color(226, 232, 240)
    pdf.line(margin, pdf.get_y(), page_w - margin, pdf.get_y())
    pdf.ln(3)
    pdf.set_font(_font, "B", 9)
    pdf.set_text_color(15, 23, 42)
    pdf.cell(0, 6, _conv_sanitize(L["disclaimer_heading"]), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(_font, "", 8.5)
    pdf.set_text_color(71, 85, 105)
    pdf.multi_cell(0, 4.8, _conv_sanitize(L["transcript_disclaimer"]))

    return bytes(pdf.output())


def render_docx(doc: Dict[str, Any], labels: Optional[Dict[str, str]] = None) -> bytes:
    """Render a structured document to DOCX bytes.

    `labels` overrides chrome strings (section headers, "Generated:")
    for the target language. When None, English defaults are used.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    L = {**DEFAULT_LABELS, **(labels or {})}

    document = Document()
    style = document.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # Title
    title = document.add_heading(doc.get("title", "Health Document"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(doc.get("subtitle", ""))
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.size = Pt(11)

    # Date
    date_p = document.add_paragraph()
    date_r = date_p.add_run(f"{L['generated']}: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    date_r.font.size = Pt(9)
    date_r.italic = True
    document.add_paragraph("_" * 60)

    # Sections
    for section in doc.get("sections", []):
        document.add_heading(section.get("heading", ""), level=1)
        document.add_paragraph(section.get("content", ""))

    # Recommendations
    recs = doc.get("recommendations", [])
    if recs:
        document.add_heading(L["recommendations"], level=1)
        for i, rec in enumerate(recs, 1):
            document.add_paragraph(f"{i}. {rec}")

    # Follow-up
    fu = doc.get("follow_up", "")
    if fu:
        document.add_heading(L["follow_up"], level=1)
        document.add_paragraph(fu)

    # Disclaimer
    document.add_paragraph("_" * 60)
    disc_p = document.add_paragraph()
    disc_r = disc_p.add_run(doc.get("disclaimer", "Generated by AMINA AI."))
    disc_r.font.size = Pt(8)
    disc_r.italic = True
    disc_r.font.color.rgb = RGBColor(120, 120, 120)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
